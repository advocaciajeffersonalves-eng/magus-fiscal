"""
MAGUS Fiscal — Módulo: Gerador de Contratos v2
Experiência inteligente com dois modos: Qualificado e Completo.
"""

import io, os, tempfile, anthropic
import streamlit as st
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ─── DADOS ──────────────────────────────────────────────────────────────────

ESTADOS_BR = [
    "Acre (AC)","Alagoas (AL)","Amapá (AP)","Amazonas (AM)","Bahia (BA)",
    "Ceará (CE)","Distrito Federal (DF)","Espírito Santo (ES)","Goiás (GO)",
    "Maranhão (MA)","Mato Grosso (MT)","Mato Grosso do Sul (MS)",
    "Minas Gerais (MG)","Pará (PA)","Paraíba (PB)","Paraná (PR)",
    "Pernambuco (PE)","Piauí (PI)","Rio de Janeiro (RJ)",
    "Rio Grande do Norte (RN)","Rio Grande do Sul (RS)","Rondônia (RO)",
    "Roraima (RR)","Santa Catarina (SC)","São Paulo (SP)","Sergipe (SE)",
    "Tocantins (TO)",
]
ESTADOS_EUA = [
    "Alabama","Alaska","Arizona","Arkansas","California","Colorado",
    "Connecticut","Delaware","Florida","Georgia","Hawaii","Idaho",
    "Illinois","Indiana","Iowa","Kansas","Kentucky","Louisiana",
    "Maine","Maryland","Massachusetts","Michigan","Minnesota",
    "Mississippi","Missouri","Montana","Nebraska","Nevada",
    "New Hampshire","New Jersey","New Mexico","New York",
    "North Carolina","North Dakota","Ohio","Oklahoma","Oregon",
    "Pennsylvania","Rhode Island","South Carolina","South Dakota",
    "Tennessee","Texas","Utah","Vermont","Virginia",
    "Washington","West Virginia","Wisconsin","Wyoming",
]
TIPOS_CONTRATO = {
    "📋  Prestação de Serviços":        "Contrato entre contratante e prestador — PJ ou PF.",
    "🤝  Parceria Comercial":           "Colaboração entre empresas ou sócios estratégicos.",
    "🔒  NDA — Confidencialidade":      "Proteção de informações sigilosas entre as partes.",
    "💰  Compra e Venda":               "Transferência de bem, produto ou ativo.",
    "👔  Contrato de Trabalho (CLT)":   "Vínculo empregatício — exclusivo para o Brasil.",
    "🌐  Independent Contractor (EUA)": "Prestador autônomo sob lei americana.",
    "🏢  Acordo de Sócios":             "Direitos e obrigações entre sócios.",
    "🏛️  Holding Familiar":             "Estrutura de holding para proteção patrimonial e sucessória.",
}
AVISO = (
    "⚠️ **Aviso Jurídico:** Minuta gerada por IA com caráter informativo e preliminar. "
    "Não substitui advogado inscrito na OAB ou *attorney* licenciado. Revise com profissional antes de assinar."
)

# ─── HELPERS ────────────────────────────────────────────────────────────────

def _claude():
    key = os.environ.get("ANTHROPIC_API_KEY","")
    return anthropic.Anthropic(api_key=key)

def _whisper(audio_bytes: bytes, filename: str = "audio.wav") -> str:
    from openai import OpenAI
    key = os.environ.get("OPENAI_API_KEY","")
    if not key:
        return "[Configure OPENAI_API_KEY para usar transcrição de voz]"
    client = OpenAI(api_key=key)
    with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as f:
        f.write(audio_bytes); tmp = f.name
    try:
        with open(tmp,"rb") as f:
            r = client.audio.transcriptions.create(model="whisper-1", file=f, language="pt")
        return r.text
    finally:
        os.unlink(tmp)

def _init():
    defaults = {
        "ct_fase": "modo",          # modo | tipo | coleta | perguntas | resultado
        "ct_modo": None,            # qualificado | completo
        "ct_tipo": None,
        "ct_jur":  "Brasil",
        "ct_estado": "",
        "ct_conversa": [],          # histórico de perguntas/respostas (modo completo)
        "ct_dados": {},             # dados coletados
        "ct_minuta": None,
        "ct_descricao_inicial": "",
        "ct_audio_gravacoes": [],   # transcrições acumuladas
        "ct_audio_contador": 0,     # chave única p/ cada gravação
    }
    for k,v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

def _reset():
    for k in list(st.session_state.keys()):
        if k.startswith("ct_"):
            del st.session_state[k]
    _init()

# ─── EXPORTAÇÃO DOCX ────────────────────────────────────────────────────────

def _gerar_docx(minuta: str, tipo: str, jurisdicao: str, estado: str) -> bytes:
    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Inches(1.2)
    sec.top_margin  = sec.bottom_margin = Inches(1.0)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MAGUS FISCAL — GERADOR DE CONTRATOS")
    r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x0A,0x1A,0x4A)

    doc.add_paragraph("─"*70)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(tipo.split("  ")[-1].upper())
    r.bold = True; r.font.size = Pt(14); r.font.color.rgb = RGBColor(0x0A,0x1A,0x4A)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Jurisdição: {jurisdicao}{f' — {estado}' if estado else ''}")
    r.italic = True; r.font.size = Pt(10)

    doc.add_paragraph("")
    p = doc.add_paragraph()
    r = p.add_run("AVISO: Minuta gerada por IA — caráter informativo e preliminar. "
                  "Não substitui aconselhamento jurídico. Revise com profissional habilitado.")
    r.font.size = Pt(8); r.italic = True; r.font.color.rgb = RGBColor(0x80,0x00,0x00)
    doc.add_paragraph("")

    for linha in minuta.split("\n"):
        p = doc.add_paragraph(); l = linha.strip()
        if l.startswith("CLÁUSULA") or l.startswith("CLAUSE") or (l.startswith("**") and l.endswith("**")):
            r = p.add_run(l.replace("**","")); r.bold = True; r.font.size = Pt(11)
        elif l.startswith("⚠️") or "PONTOS DE ATENÇÃO" in l:
            r = p.add_run(l); r.bold = True; r.font.color.rgb = RGBColor(0xCC,0x44,0x00)
        else:
            p.add_run(l)
        p.paragraph_format.space_after = Pt(4)

    doc.add_paragraph("")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Gerado por MAGUS Fiscal em {datetime.now().strftime('%d/%m/%Y às %H:%M')}")
    r.font.size = Pt(8); r.italic = True; r.font.color.rgb = RGBColor(0x80,0x80,0x80)

    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.read()

# ─── PROMPTS ────────────────────────────────────────────────────────────────

def _system_gerador(tipo: str, jur: str, estado: str, modo: str) -> str:
    if jur == "Brasil":
        base = f"Jurisdição: Brasil — Direito Civil. Estado: {estado}. Código Civil, CLT, LGPD quando aplicável."
    elif jur == "Estados Unidos":
        base = f"Jurisdição: EUA — {estado} (Common Law Estadual). Elementos: offer, acceptance, consideration, capacity, lawful purpose."
    else:
        base = f"Jurisdição: Internacional Brasil/EUA. Incluir: lei aplicável, foro, idioma prevalente, moeda, tributos."

    qualif_instr = (
        "Gere contrato com [DADO A COMPLETAR] nos campos não fornecidos. "
        "No final, liste em 'CAMPOS PARA COMPLETAR' todos os dados que o usuário precisa preencher."
    ) if modo == "qualificado" else (
        "Use TODOS os dados fornecidos. Gere contrato completo e pronto para assinar, sem lacunas."
    )

    return f"""Você é o MAGUS Contratos — especialista em minutas contratuais para empresários brasileiros.

{base}
Tipo: {tipo}
Modo: {'QUALIFICADO (com lacunas indicadas)' if modo=='qualificado' else 'COMPLETO (pronto para assinar)'}

REGRAS:
1. {qualif_instr}
2. Estrutura: cabeçalho → qualificação → CLÁUSULA PRIMEIRA, SEGUNDA... → assinaturas → PONTOS DE ATENÇÃO
3. Marque cláusulas de alto risco com ⚠️
4. Nunca afirme que o contrato é "válido" ou "garantido"
5. Sempre recomende revisão por advogado habilitado
6. Responda em português do Brasil"""

def _system_entrevistador(tipo: str, jur: str, estado: str) -> str:
    return f"""Você é um assistente especialista em coleta de informações para elaboração de contratos.

Tipo de contrato: {tipo}
Jurisdição: {jur}{f' — {estado}' if estado else ''}

Sua missão: coletar todas as informações necessárias para gerar um contrato COMPLETO e pronto para assinar.

REGRAS DA ENTREVISTA:
1. Analise o que o usuário já informou
2. Faça UMA pergunta por vez — a mais importante que ainda falta
3. Seja direto e objetivo — sem rodeios
4. Quando tiver TODAS as informações essenciais, responda com:
   "PRONTO ✅ Tenho todas as informações. Posso gerar seu contrato agora."
5. Organize as perguntas por prioridade: partes → objeto → valores → prazos → cláusulas especiais
6. Responda sempre em português do Brasil

Informações essenciais para este contrato:
- Identificação completa das partes (nome, documento, endereço, representante)
- Objeto/serviço detalhado
- Valores e forma de pagamento
- Prazo e condições de rescisão
- Dados específicos do tipo de contrato"""

# ─── TELA 1: ESCOLHA DO MODO ────────────────────────────────────────────────

def _tela_modo():
    st.markdown("""
    <div style='background:linear-gradient(90deg,#0a1a4a,#1a3a7a);padding:20px 24px;
    border-radius:10px;margin-bottom:24px'>
    <h2 style='color:white;margin:0'>📄 Gerador de Contratos</h2>
    <p style='color:#a0c0ff;margin:6px 0 0'>Minutas profissionais para Brasil e Estados Unidos</p>
    </div>""", unsafe_allow_html=True)

    st.warning(AVISO)
    st.markdown("### Como você quer receber seu contrato?")
    st.markdown("")

    col1, col2 = st.columns(2)

    with col1:
        st.markdown("""
        <div style='border:2px solid #4a7acc;border-radius:12px;padding:20px;min-height:180px'>
        <h3 style='color:#7ab0ff;margin-top:0'>📋 Qualificado</h3>
        <p style='color:#ccd8f0'>Receba a minuta <b>agora</b> com espaços para preencher.<br><br>
        Ideal para quem quer um modelo base e vai revisar com advogado.</p>
        <ul style='color:#a0b8d8;font-size:0.9rem'>
        <li>Rápido — gerado em segundos</li>
        <li>Campos indicados para completar</li>
        <li>Lista do que falta para mais segurança</li>
        </ul></div>""", unsafe_allow_html=True)
        st.markdown("")
        if st.button("📋  Quero modelo com lacunas", use_container_width=True,
                     key="btn_modo_qualif", type="primary"):
            st.session_state.ct_modo = "qualificado"
            st.session_state.ct_fase = "tipo"
            st.rerun()

    with col2:
        st.markdown("""
        <div style='border:2px solid #c8973a;border-radius:12px;padding:20px;min-height:180px'>
        <h3 style='color:#e8b84a;margin-top:0'>✅ Completo</h3>
        <p style='color:#f0e0c0'>Eu coletarei <b>todas as informações</b> necessárias e entregarei o contrato pronto para assinar.<br><br>
        Ideal para quem quer o documento final.</p>
        <ul style='color:#d0b888;font-size:0.9rem'>
        <li>Sem lacunas — pronto para assinar</li>
        <li>Coleta por texto, documento ou voz</li>
        <li>Perguntas inteligentes até completar</li>
        </ul></div>""", unsafe_allow_html=True)
        st.markdown("")
        if st.button("✅  Quero contrato completo", use_container_width=True,
                     key="btn_modo_completo", type="primary"):
            st.session_state.ct_modo = "completo"
            st.session_state.ct_fase = "tipo"
            st.rerun()

# ─── TELA 2: TIPO E JURISDIÇÃO ──────────────────────────────────────────────

def _tela_tipo():
    modo = st.session_state.ct_modo
    badge = "📋 Qualificado" if modo == "qualificado" else "✅ Completo"
    cor   = "#1a3a7a" if modo == "qualificado" else "#c8973a"

    st.markdown(f"""
    <div style='background:linear-gradient(90deg,#0a1a4a,#1a3a7a);padding:16px 24px;
    border-radius:10px;margin-bottom:20px;display:flex;align-items:center;gap:12px'>
    <h2 style='color:white;margin:0'>📄 Gerador de Contratos</h2>
    <span style='background:{cor};color:white;padding:3px 12px;border-radius:20px;
    font-size:0.8rem;font-weight:700'>{badge}</span>
    </div>""", unsafe_allow_html=True)

    col_back, _ = st.columns([1,4])
    with col_back:
        if st.button("← Voltar", key="back_tipo"):
            st.session_state.ct_fase = "modo"; st.rerun()

    st.markdown("### 1. Tipo de Contrato")
    tipo = st.selectbox("Selecione:", list(TIPOS_CONTRATO.keys()),
                        key="sel_tipo_contrato",
                        index=list(TIPOS_CONTRATO.keys()).index(st.session_state.ct_tipo)
                        if st.session_state.ct_tipo else 0)
    st.caption(f"_{TIPOS_CONTRATO[tipo]}_")

    st.markdown("---")
    st.markdown("### 2. Jurisdição")

    # Restrições de tipo
    if tipo == "👔  Contrato de Trabalho (CLT)":
        jur = "Brasil"
        st.info("ℹ️ Contrato de Trabalho (CLT) é exclusivo para a legislação brasileira.")
        estado = st.selectbox("Estado (UF) do vínculo:", ESTADOS_BR, key="estado_clt")
    elif tipo == "🌐  Independent Contractor (EUA)":
        jur = "Estados Unidos"
        st.info("ℹ️ Independent Contractor é exclusivo para a legislação americana.")
        estado = st.selectbox("Estado americano (obrigatório):", ESTADOS_EUA, key="estado_ic")
    else:
        col1, col2, col3 = st.columns(3)
        jur_atual = st.session_state.get("ct_jur","Brasil")

        with col1:
            if st.button("🇧🇷  Brasil", use_container_width=True, key="jur_br",
                         type="primary" if jur_atual=="Brasil" else "secondary"):
                st.session_state.ct_jur = "Brasil"; st.rerun()
        with col2:
            if st.button("🇺🇸  Estados Unidos", use_container_width=True, key="jur_eua",
                         type="primary" if jur_atual=="Estados Unidos" else "secondary"):
                st.session_state.ct_jur = "Estados Unidos"; st.rerun()
        with col3:
            if st.button("🌐  Internacional", use_container_width=True, key="jur_int",
                         type="primary" if jur_atual=="Internacional" else "secondary"):
                st.session_state.ct_jur = "Internacional"; st.rerun()

        jur = st.session_state.ct_jur

        # Explicação contextual da jurisdição escolhida
        if jur == "Brasil":
            st.success("🇧🇷 **Brasil:** contrato regido pelo Código Civil brasileiro. Escolha o estado de referência para definir o foro.")
            estado = st.selectbox("Estado (UF) de referência:", ESTADOS_BR, key="estado_br")
        elif jur == "Estados Unidos":
            st.warning("🇺🇸 **Estados Unidos:** as regras contratuais variam por estado. Escolha o estado cuja lei vai reger o contrato.")
            estado = st.selectbox("Estado americano (obrigatório):", ESTADOS_EUA, key="estado_eua")
        else:
            st.info("""🌐 **Internacional:** para contratos com partes em dois países — por exemplo, uma empresa brasileira contratando com empresa americana.

O contrato precisará definir: **qual lei rege** (brasileira ou americana), **onde resolver disputas** (foro ou arbitragem), **qual idioma prevalece** e **qual moeda** é usada.""")
            col_i1, col_i2 = st.columns(2)
            with col_i1:
                estado_br  = st.selectbox("Estado brasileiro:", ESTADOS_BR, key="estado_br_int")
            with col_i2:
                estado_eua = st.selectbox("Estado americano:", ESTADOS_EUA, key="estado_eua_int")
            estado = f"{estado_br} (BR) / {estado_eua} (EUA)"

    st.markdown("---")
    if st.button("Continuar →", type="primary", use_container_width=False, key="btn_avancar_tipo"):
        st.session_state.ct_tipo   = tipo
        st.session_state.ct_jur    = jur
        st.session_state.ct_estado = estado
        st.session_state.ct_fase   = "coleta"
        st.rerun()

# ─── ABA DE ÁUDIO INTELIGENTE (compartilhada) ────────────────────────────────

def _aba_audio(modo: str):
    """
    Gravação com acúmulo:
    - Cada gravação é transcrita automaticamente ao parar
    - Transcrições se acumulam em ct_audio_gravacoes
    - Exibe "É isso mesmo?" antes de prosseguir
    - Botão "Adicionar mais áudio" permite N gravações
    """
    gravacoes = st.session_state.ct_audio_gravacoes

    # ── Área de acúmulo ──────────────────────────────────────────────────
    if gravacoes:
        st.markdown("**📝 O que você disse até agora:**")
        texto_acumulado = "\n\n".join(gravacoes)
        texto_final = st.text_area(
            "Revise e edite se necessário:",
            value=texto_acumulado,
            key="ct_audio_acumulado",
            height=150,
        )

        col1, col2 = st.columns(2)
        with col1:
            label_btn = "📄 Gerar Minuta com Lacunas →" if modo == "qualificado" else "✅ É isso — continuar →"
            if st.button(label_btn, type="primary", key="btn_confirmar_audio",
                         use_container_width=True):
                st.session_state.ct_audio_gravacoes = []
                st.session_state.ct_audio_contador  = 0
                st.session_state.ct_descricao_inicial = texto_final
                if modo == "qualificado":
                    st.session_state.ct_dados["objeto"] = texto_final
                    _gerar_e_mostrar("qualificado")
                else:
                    _iniciar_entrevista(texto_final)
                return
        with col2:
            if st.button("🎙️ Adicionar mais áudio", key="btn_add_audio",
                         use_container_width=True):
                st.session_state.ct_audio_contador += 1
                st.rerun()

        st.markdown("---")
        st.caption("Ou grave um trecho complementar abaixo:")
    else:
        st.markdown("**🎙️ Grave sua explicação — transcrevo automaticamente ao parar:**")
        st.caption(
            "Fale tudo que quiser. Pode parar e gravar de novo quantas vezes precisar — "
            "as transcrições se acumulam até você confirmar."
        )

    # ── Gravador com chave única por rodada ──────────────────────────────
    contador = st.session_state.ct_audio_contador
    audio = st.audio_input(
        "🎙️ Clique para gravar" if not gravacoes else "🎙️ Gravar complemento:",
        key=f"ct_audio_{contador}",
    )

    if audio:
        with st.spinner("🎙️ Transcrevendo..."):
            texto = _whisper(audio.read())
        if texto and not texto.startswith("[Configure"):
            st.session_state.ct_audio_gravacoes.append(texto)
            st.session_state.ct_audio_contador += 1   # próxima gravação = nova chave
            st.rerun()
        else:
            st.error(texto or "Não foi possível transcrever. Verifique sua chave OpenAI.")


# ─── TELA 3A: COLETA — MODO QUALIFICADO ─────────────────────────────────────

def _tela_coleta_qualificado():
    tipo   = st.session_state.ct_tipo
    jur    = st.session_state.ct_jur
    estado = st.session_state.ct_estado

    col_back, _ = st.columns([1,4])
    with col_back:
        if st.button("← Voltar", key="back_qualif"):
            st.session_state.ct_audio_gravacoes = []
            st.session_state.ct_audio_contador  = 0
            st.session_state.ct_fase = "tipo"
            st.rerun()

    st.markdown(f"### 📋 Descreva seu caso — {tipo.split('  ')[-1]}")
    st.markdown(f"**Jurisdição:** {jur}{f' — {estado}' if estado else ''}")
    st.info("📋 **Minuta com lacunas:** descreva o que tiver — texto, voz ou arquivo. "
            "O que não informar ficará como **[DADO A COMPLETAR]** na minuta.")

    tab_texto, tab_audio, tab_arquivo = st.tabs(["✍️ Digitar", "🎙️ Falar", "📎 Enviar arquivo"])

    with tab_texto:
        texto = st.text_area(
            "Descreva seu caso:",
            value=st.session_state.get("ct_descricao_inicial", ""),
            placeholder=(
                "Exemplo: Preciso de um contrato de prestação de serviços. "
                "Sou advogado tributarista e vou prestar consultoria para empresa ABC Ltda. "
                "Pode deixar os dados específicos em branco — só preciso do modelo base."
            ),
            key="ct_texto_qualif", height=160,
        )
        if st.button("📄 Gerar Minuta com Lacunas →", type="primary",
                     key="btn_gerar_qualif_texto", disabled=not texto.strip()):
            st.session_state.ct_descricao_inicial = texto
            st.session_state.ct_dados["objeto"]   = texto
            _gerar_e_mostrar("qualificado")

    with tab_audio:
        _aba_audio("qualificado")

    with tab_arquivo:
        st.markdown("**Envie um documento com as informações** (PDF, DOCX ou TXT):")
        arq = st.file_uploader("Selecione o arquivo:", type=["pdf","docx","txt"],
                               key="ct_arquivo_qualif")
        if arq:
            texto_arq = _extrair_arquivo(arq)
            if texto_arq:
                st.success(f"✅ {arq.name} lido — {len(texto_arq)} caracteres")
                texto_edit = st.text_area("Conteúdo extraído (edite se necessário):",
                                          value=texto_arq[:3000], key="ct_texto_arq_qualif",
                                          height=140)
                if st.button("📄 Gerar Minuta com Lacunas →", type="primary",
                             key="btn_gerar_qualif_arq"):
                    st.session_state.ct_dados["objeto"] = texto_edit
                    _gerar_e_mostrar("qualificado")

# ─── TELA 3B: COLETA — MODO COMPLETO ────────────────────────────────────────

def _tela_coleta_completo():
    tipo   = st.session_state.ct_tipo
    jur    = st.session_state.ct_jur
    estado = st.session_state.ct_estado

    col_back, _ = st.columns([1,4])
    with col_back:
        if st.button("← Voltar", key="back_completo"):
            st.session_state.ct_audio_gravacoes = []
            st.session_state.ct_audio_contador  = 0
            st.session_state.ct_fase = "tipo"
            st.rerun()

    st.markdown(f"### ✅ Conte seu caso — {tipo.split('  ')[-1]}")
    st.markdown(f"**Jurisdição:** {jur}{f' — {estado}' if estado else ''}")
    st.markdown("Descreva sua situação da forma que preferir. Quanto mais detalhes, melhor o contrato gerado.")

    # ── Abas de entrada ──
    tab_texto, tab_audio, tab_arquivo = st.tabs(["✍️ Digitar", "🎙️ Falar", "📎 Enviar arquivo"])

    texto_coletado = ""

    with tab_texto:
        texto_coletado = st.text_area(
            "Descreva seu caso:",
            value=st.session_state.get("ct_descricao_inicial",""),
            placeholder=(
                "Exemplo: Sou advogado em Goiás e preciso de um contrato de prestação de serviços. "
                "Vou prestar consultoria tributária para a empresa ABC Ltda., CNPJ 00.000.000/0001-00, "
                "localizada em Goiânia. Os honorários são R$ 5.000 por mês, pagamento todo dia 5. "
                "Prazo de 12 meses, podendo ser renovado..."
            ),
            key="ct_texto_inicial", height=160
        )
        if st.button("Enviar descrição →", key="btn_enviar_texto",
                     type="primary", disabled=not texto_coletado.strip()):
            st.session_state.ct_descricao_inicial = texto_coletado
            _iniciar_entrevista(texto_coletado)

    with tab_audio:
        _aba_audio("completo")

    with tab_arquivo:
        st.markdown("**Envie um documento com as informações** (PDF, DOCX ou TXT):")
        arq = st.file_uploader("Selecione o arquivo:", type=["pdf","docx","txt"],
                                key="ct_arquivo")
        if arq:
            texto_arq = _extrair_arquivo(arq)
            if texto_arq:
                st.success(f"✅ {arq.name} lido — {len(texto_arq)} caracteres")
                texto_edit = st.text_area("Conteúdo extraído (edite se necessário):",
                                           value=texto_arq[:3000], key="ct_texto_arq", height=140)
                if st.button("Usar este documento →", key="btn_usar_arq", type="primary"):
                    st.session_state.ct_descricao_inicial = texto_edit
                    _iniciar_entrevista(texto_edit)

# ─── ENTREVISTA INTELIGENTE ──────────────────────────────────────────────────

def _iniciar_entrevista(descricao_inicial: str):
    """Começa o processo de Q&A com Claude."""
    st.session_state.ct_conversa = [
        {"role": "user", "content": descricao_inicial}
    ]
    st.session_state.ct_fase = "perguntas"
    st.rerun()

def _tela_perguntas():
    tipo   = st.session_state.ct_tipo
    jur    = st.session_state.ct_jur
    estado = st.session_state.ct_estado
    conversa = st.session_state.ct_conversa

    col_back, _ = st.columns([1,4])
    with col_back:
        if st.button("← Recomeçar", key="back_perguntas"):
            st.session_state.ct_fase = "coleta"; st.rerun()

    st.markdown(f"### ✅ Coletando informações — {tipo.split('  ')[-1]}")
    st.progress(min(len(conversa) / 14, 0.95), text="Coletando dados...")

    # Exibir histórico da conversa
    for msg in conversa:
        if msg["role"] == "user":
            with st.chat_message("user", avatar="👤"):
                st.write(msg["content"])
        else:
            with st.chat_message("assistant", avatar="⚖️"):
                st.write(msg["content"])
                # Verificar se Claude sinalizou que tem tudo
                if "PRONTO ✅" in msg["content"]:
                    st.markdown("")
                    if st.button("📄  Gerar Contrato Completo Agora", type="primary",
                                 use_container_width=True, key="btn_gerar_completo"):
                        _gerar_e_mostrar("completo")
                    return

    # Se último msg é do usuário, pedir resposta do Claude
    if not conversa or conversa[-1]["role"] == "user":
        with st.spinner("⚖️ Analisando..."):
            sys = _system_entrevistador(tipo, jur, estado)
            resp = _claude().messages.create(
                model="claude-sonnet-4-6",
                max_tokens=600,
                system=sys,
                messages=conversa
            )
            resposta = resp.content[0].text
            st.session_state.ct_conversa.append({"role": "assistant", "content": resposta})
            st.rerun()

    # Campo de resposta do usuário
    else:
        # Input por texto
        resposta_user = st.chat_input("Sua resposta...", key="chat_resposta")
        if resposta_user:
            st.session_state.ct_conversa.append({"role": "user", "content": resposta_user})
            st.rerun()

        # Opção de responder por voz — auto-adiciona ao chat após transcrição
        with st.expander("🎙️ Responder por voz"):
            st.caption("Grave sua resposta. Ao parar, transcrevo e adiciono automaticamente.")
            audio_resp = st.audio_input("Grave:", key=f"audio_resp_{len(conversa)}")
            if audio_resp:
                with st.spinner("🎙️ Transcrevendo..."):
                    texto = _whisper(audio_resp.read())
                if texto and not texto.startswith("[Configure"):
                    st.info(f"📝 Entendi: _{texto}_")
                    st.session_state.ct_conversa.append({"role": "user", "content": texto})
                    st.rerun()
                else:
                    st.error(texto or "Erro na transcrição.")

# ─── GERAÇÃO DO CONTRATO ─────────────────────────────────────────────────────

def _gerar_e_mostrar(modo: str):
    tipo   = st.session_state.ct_tipo
    jur    = st.session_state.ct_jur
    estado = st.session_state.ct_estado

    # Montar contexto completo
    if modo == "qualificado":
        dados = st.session_state.ct_dados
        contexto = "\n".join(f"{k}: {v}" for k,v in dados.items() if v)
        if not contexto:
            contexto = "Gerar modelo genérico com todos os campos a completar."
    else:
        # Usar toda a conversa coletada
        contexto = "INFORMAÇÕES COLETADAS NA ENTREVISTA:\n\n"
        for msg in st.session_state.ct_conversa:
            papel = "USUÁRIO" if msg["role"] == "user" else "ASSISTENTE"
            contexto += f"[{papel}]: {msg['content']}\n\n"

    system = _system_gerador(tipo, jur, estado, modo)
    prompt = (
        f"Tipo de contrato: {tipo}\n"
        f"Jurisdição: {jur}{f' — {estado}' if estado else ''}\n"
        f"Modo: {modo.upper()}\n"
        f"Data de referência: {datetime.now().strftime('%d/%m/%Y')}\n\n"
        f"{contexto}\n\n"
        f"Gere a minuta contratual completa agora."
    )

    with st.spinner("📝 Gerando minuta contratual..."):
        resp = _claude().messages.create(
            model="claude-opus-4-7",
            max_tokens=4096,
            system=system,
            messages=[{"role": "user", "content": prompt}]
        )
        minuta = resp.content[0].text

    st.session_state.ct_minuta = minuta
    st.session_state.ct_fase   = "resultado"
    st.rerun()

# ─── TELA RESULTADO ──────────────────────────────────────────────────────────

def _tela_resultado():
    tipo   = st.session_state.ct_tipo
    jur    = st.session_state.ct_jur
    estado = st.session_state.ct_estado
    minuta = st.session_state.ct_minuta
    modo   = st.session_state.ct_modo

    st.markdown(f"### 📄 Minuta Gerada — {tipo.split('  ')[-1]}")
    st.warning(AVISO)

    badge = "📋 Qualificado (com lacunas)" if modo == "qualificado" else "✅ Completo"
    st.caption(f"{badge} · {jur}{f' — {estado}' if estado else ''} · {datetime.now().strftime('%d/%m/%Y %H:%M')}")

    with st.expander("👁️ Visualizar minuta completa", expanded=True):
        st.markdown(minuta)

    st.markdown("#### ⬇️ Exportar")
    col1, col2, col3 = st.columns(3)

    nome = (f"MAGUS_{tipo.split('  ')[-1].replace(' ','_')}_"
            f"{jur.replace(' ','_')}_{datetime.now().strftime('%Y%m%d_%H%M')}")

    with col1:
        docx = _gerar_docx(minuta, tipo, jur, estado)
        st.download_button("📥 Baixar DOCX", data=docx,
                           file_name=f"{nome}.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                           use_container_width=True)
    with col2:
        st.download_button("📥 Baixar TXT", data=minuta.encode("utf-8"),
                           file_name=f"{nome}.txt", mime="text/plain",
                           use_container_width=True)
    with col3:
        if st.button("🔄 Novo contrato", use_container_width=True, key="btn_novo"):
            _reset(); st.rerun()

# ─── EXTRAÇÃO DE ARQUIVO ─────────────────────────────────────────────────────

def _extrair_arquivo(arq) -> str:
    try:
        if arq.name.endswith(".txt"):
            return arq.read().decode("utf-8", errors="ignore")
        elif arq.name.endswith(".pdf"):
            import pdfplumber
            with pdfplumber.open(arq) as pdf:
                return "\n".join(p.extract_text() or "" for p in pdf.pages)
        elif arq.name.endswith(".docx"):
            from docx import Document as DocxDoc
            doc = DocxDoc(arq)
            return "\n".join(p.text for p in doc.paragraphs)
    except Exception as e:
        st.error(f"Erro ao ler arquivo: {e}")
    return ""

# ─── RENDER PRINCIPAL ────────────────────────────────────────────────────────

def render_contratos():
    _init()

    fase = st.session_state.ct_fase
    modo = st.session_state.ct_modo

    if fase == "modo":
        _tela_modo()
    elif fase == "tipo":
        _tela_tipo()
    elif fase == "coleta":
        if modo == "qualificado":
            _tela_coleta_qualificado()
        else:
            _tela_coleta_completo()
    elif fase == "perguntas":
        _tela_perguntas()
    elif fase == "resultado":
        _tela_resultado()
